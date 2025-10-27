"""
Document Parser Utility
Supports parsing PDF, DOCX, DOC files using MinerU and python-docx
"""

import os
import tempfile
from typing import Optional, Dict
from pathlib import Path

from utils.logger import logger

try:
    from magic_pdf.data.dataset import PymuDocDataset
    MINERU_AVAILABLE = True
except ImportError as e:
    MINERU_AVAILABLE = False
    logger.warning(f"MinerU not available: {e}")

try:
    import fitz  # type: ignore[attr-defined]
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF (fitz) not available; PDF text extraction may be limited")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available")


class DocumentParser:
    """Parse various document formats to extract text content"""
    
    def __init__(self):
        self.temp_dir = tempfile.mkdtemp(prefix="youtu_graphrag_")
        logger.info(f"DocumentParser initialized with temp dir: {self.temp_dir}")
    
    def parse_file(self, file_path: str, file_type: str) -> Optional[str]:
        """
        Parse a document file and extract text content
        
        Args:
            file_path: Path to the document file
            file_type: File extension (.pdf, .docx, .doc)
            
        Returns:
            Extracted text content or None if parsing fails
        """
        file_type = file_type.lower()
        
        try:
            if file_type == '.pdf':
                return self._parse_pdf(file_path)
            elif file_type in ['.docx', '.doc']:
                return self._parse_docx(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_type}")
                return None
        except Exception as e:
            logger.error(f"Error parsing {file_type} file: {e}")
            return None
    
    def _parse_pdf(self, pdf_path: str) -> Optional[str]:
        """
        Parse PDF using MinerU
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Extracted text content
        """
        if not MINERU_AVAILABLE:
            logger.error("MinerU is not installed. Cannot parse PDF files.")
            return None
        
        try:
            with open(pdf_path, 'rb') as f:
                pdf_bytes = f.read()
        except Exception as e:
            logger.error(f"Unable to read PDF file {pdf_path}: {e}")
            return None

        # -------- MinerU / PyMuPDF pipeline --------
        if MINERU_AVAILABLE:
            try:
                dataset = PymuDocDataset(pdf_bytes, lang='auto')
                text_parts: list[str] = []
                for page_index in range(len(dataset)):
                    try:
                        page_doc = dataset.get_page(page_index).get_doc()
                        page_text = page_doc.get_text("text")
                        if page_text:
                            text_parts.append(page_text.strip())
                    except Exception as page_err:
                        logger.warning(
                            f"MinerU pipeline failed to extract page {page_index} of {pdf_path}: {page_err}"
                        )
                        continue

                if text_parts:
                    extracted_text = '\n\n'.join(text_parts)
                    logger.info(
                        f"Successfully extracted {len(extracted_text)} chars from PDF via MinerU pipeline"
                    )
                    return extracted_text
            except Exception as e:
                logger.error(f"MinerU PDF parsing failed: {e}")

        # -------- Direct PyMuPDF fallback --------
        if PYMUPDF_AVAILABLE:
            try:
                text_parts: list[str] = []
                with fitz.open(pdf_path) as doc:  # type: ignore[attr-defined]
                    for page in doc:
                        page_text = page.get_text("text")
                        if page_text:
                            text_parts.append(page_text.strip())
                if text_parts:
                    extracted_text = '\n\n'.join(text_parts)
                    logger.info(
                        f"Successfully extracted {len(extracted_text)} chars from PDF via PyMuPDF fallback"
                    )
                    return extracted_text
            except Exception as e:
                logger.error(f"PyMuPDF fallback failed: {e}")

        # -------- PyPDF fallback --------
        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(pdf_path)
            text_parts = []
            for page_index, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                except Exception as page_err:
                    logger.warning(
                        f"pypdf failed to extract page {page_index} of {pdf_path}: {page_err}"
                    )
                    continue
                if page_text:
                    text_parts.append(page_text.strip())

            if text_parts:
                extracted_text = '\n\n'.join(text_parts)
                logger.info(
                    f"Successfully extracted {len(extracted_text)} chars from PDF via pypdf fallback"
                )
                return extracted_text
        except ImportError:
            logger.debug("pypdf not installed; skipping pypdf fallback")
        except Exception as e:
            logger.error(f"pypdf fallback failed: {e}")

        logger.error(f"Unable to extract text from PDF: {pdf_path}")
        return None
    
    def _parse_docx(self, docx_path: str) -> Optional[str]:
        """
        Parse DOCX/DOC using python-docx
        
        Args:
            docx_path: Path to DOCX/DOC file
            
        Returns:
            Extracted text content
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx is not installed. Cannot parse DOCX/DOC files.")
            return None
        
        try:
            doc = DocxDocument(docx_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            extracted_text = '\n'.join(text_parts)
            
            if not extracted_text.strip():
                logger.warning(f"No text extracted from DOCX: {docx_path}")
                return None
            
            logger.info(f"Successfully extracted {len(extracted_text)} chars from DOCX")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            return None
    
    def cleanup(self):
        """Clean up temporary files"""
        try:
            import shutil
            if os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
                logger.info(f"Cleaned up temp dir: {self.temp_dir}")
        except Exception as e:
            logger.error(f"Error cleaning up temp dir: {e}")


# Global parser instance
_parser_instance = None

def get_parser() -> DocumentParser:
    """Get or create global parser instance"""
    global _parser_instance
    if _parser_instance is None:
        _parser_instance = DocumentParser()
    return _parser_instance
